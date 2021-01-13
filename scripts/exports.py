import sys
import itertools

import mysql.connector as mysql
from mysql.connector import Error

import pandas as pd

from openpyxl import load_workbook  # For .xlsx conversion
import pyreadstat  # For .sav and .xpt conversion
from docx import Document  # For .docx conversion

import config

def exportToCsv(name, df, attributes):
    # Open the csv file to write to
    file_name = name + '.csv'
    csv_file = open(file_name, 'w')

    # Write data frame to .CSV file
    df.to_csv(csv_file, header=True, columns=attributes)

    # Close file write stream
    csv_file.close()


def exportToSav(name, df, attributes):
    # Write data frame to sav (.SPSS) file
    file_name = name + '.sav'
    pyreadstat.write_sav(df, file_name, column_labels=attributes)


def exportToExcel(name, df):
    try:
        # Append data frame to excel (.XLSX) file
        with pd.ExcelWriter('output.xlsx', mode = 'a') as writer:
            df.to_excel(writer, sheet_name = name)

    # If excel file not created yet, make initial write to it
    except FileNotFoundError as e:
        # Write data frame to excel (.XLSX) file
        with pd.ExcelWriter('output.xlsx') as writer:
            df.to_excel(writer, sheet_name='Funder')


def exportToDocument(name, df):
    # Create document file object
    doc = Document()

    # Write data frame to word (.DOCX) file
    table = doc.add_table(df.shape[0] + 1, df.shape[1])
    table.style = 'Table Grid'

    # Add the header rows to the table
    for column in range(df.shape[1]):
        table.cell(0, column).text = df.columns[column]

    # Add the data rows to the table
    cells = table._cells
    for row in range(df.shape[0]):
        table.row_cells = cells[(row + 1) * df.shape[1]: (row + 2) * df.shape[1]]
        for column in range(df.shape[1]):
            table.row_cells[column].text = str(df.values[row, column])

    # Save the document file
    file_name = name + '.docx'
    doc.save(file_name)


def exportToXpt(name, df, attributes):
    # Write data frame to sas file transport format (.XPT)
    file_name = name + '.xpt'
    pyreadstat.write_xport(df, file_name, column_labels=attributes)

def exportToHTML(name, df):
    # Set up formatting to create a basic HTML page containing tables
    pageStart = "<!DOCTYPE html>\n<html lang='en'>\n\t<head>\n\t\t<title>\n\t\tgirlseducation to HTML\n\t\t</title>\n\t</head>\n\n\t<body>\n"
    pageEnd = "\n\t</body>\n</html>"

    page = pageStart + df.to_html() + pageEnd

    fileName = name + ".html"

    # Create and write to the new HTML file
    newFile = open(fileName, "wt")
    newFile.write(page)
    newFile.close()

def exportToRTF(name, df):
    # Set up formatting for rtf document
    page = "{\\rtf1" + df.to_string() + "}"

    fileName = name + ".rtf"

    # Create and write to the new RTF file
    newFile = open(fileName, "wt")
    newFile.write(page)
    newFile.close()

def exportToTXT(name, df):
    fileName = name + ".txt"

    # Create and write to the new RTF file
    newFile = open(fileName, "wt")
    newFile.write(df.to_string())
    newFile.close()


# Main method - accepts command line arguments
# Argument at index 0 is the program name,
# Argument at index 1 is the entity type (table name in the database from config file),
# Argument at index 2 is the file format to export to,
# Command format: python exports.py <entity> <file format>
# For example, python exports.py initiative excel
def main():
    # Open connection with database
    connection = None
    try:
        # Set up database connection - manually change username, password and db name as needed
        connection = mysql.connect(user=config.db_username, passwd=config.db_password, db=config.database)
        print("Connection to MySQL database " + config.database + " successful")
    except Error as e:
        print("The error {!r} occurred".format(e, e.args[0]))

    # If connection to database can be made,
    # Get attributes and data rows, set into a tabular data collection frame, and write to appropriate files
    if connection is not None:
        # Set up cursor and perform query for all rows of table
        cur = connection.cursor()

        try:
            #query the table names in the given database
            query = ("SELECT TABLE_NAME " 
            "FROM INFORMATION_SCHEMA.TABLES "
            "WHERE TABLE_TYPE = 'BASE TABLE' AND TABLE_SCHEMA = %(NAME)s")
                     
            cur.execute(query, {'NAME':config.database})

            table_names = list(itertools.chain(*cur.fetchall()))
            if sys.argv[1] in table_names:
                cur.execute("select * from {0}" .format(sys.argv[1]))

                # Get columns corresponding to attributes
                attributes = [col for col in cur.column_names]
                if len(attributes) > 0:
                    # Set data frame with rows
                    df = pd.DataFrame(cur.fetchall(), columns=attributes)
                    if len(df) > 0:
                        if sys.argv[2] == 'all':
                            exportToCsv(sys.argv[1], df, attributes)
                            exportToSav(sys.argv[1], df, attributes)
                            exportToExcel(sys.argv[1], df)
                            exportToDocument(sys.argv[1], df)
                            exportToXpt(sys.argv[1], df, attributes)
                            exportToHTML(sys.argv[1], df)
                            exportToRTF(sys.argv[1], df)
                            exportToTXT(sys.argv[1], df)
                            print('Successfully exported all file types for table: '+ sys.argv[1])
                        elif sys.argv[2] == 'csv':
                            exportToCsv(sys.argv[1], df, attributes)
                            print('Successfully exported csv for table: '+ sys.argv[1])
                        elif sys.argv[2] == 'sav':
                            exportToSav(sys.argv[1], df, attributes)
                            print('Successfully exported sav for table: '+ sys.argv[1])
                        elif sys.argv[2] == 'excel':
                            exportToExcel(sys.argv[1], df)
                            print('Successfully exported excel for table: '+ sys.argv[1])
                        elif sys.argv[2] == 'word':
                            exportToDocument(sys.argv[1], df)
                            print('Successfully exported word for table: '+ sys.argv[1])
                        elif sys.argv[2] == 'sas':
                            exportToXpt(sys.argv[1], df, attributes)
                            print('Successfully exported sas for table: '+ sys.argv[1])
                        elif sys.argv[2] == 'html':
                            exportToHTML(sys.argv[1], df)
                            print('Successfully exported html for table: '+ sys.argv[1])
                        elif sys.argv[2] == 'rtf':
                            exportToRTF(sys.argv[1], df)
                            print('Successfully exported rtf for table: '+ sys.argv[1])
                        elif sys.argv[2] == 'txt':
                            exportToTXT(sys.argv[1], df)
                            print('Successfully exported txt for table: '+ sys.argv[1])
                        else:
                            print('Invalid file format')
                    else:
                        print('No data is available to be retrieved')
                else:
                    print('No table attributes found')
            else:
                print(sys.argv[1] + " table not found in database: " + config.database)
        except IndexError as e:
            print("Missing argument(s)")

    # Close database connection
    cur.close()
    connection.close()

if __name__ == "__main__":
    main()
